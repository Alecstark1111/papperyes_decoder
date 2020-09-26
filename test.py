from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def list_number(doc, par, prev=None, level=None, num=True):
    """
    Makes a paragraph into a list item with a specific level and
    optional restart.

    An attempt will be made to retreive an abstract numbering style that
    corresponds to the style of the paragraph. If that is not possible,
    the default numbering or bullet style will be used based on the
    ``num`` parameter.

    Parameters
    ----------
    doc : docx.document.Document
        The document to add the list into.
    par : docx.paragraph.Paragraph
        The paragraph to turn into a list item.
    prev : docx.paragraph.Paragraph or None
        The previous paragraph in the list. If specified, the numbering
        and styles will be taken as a continuation of this paragraph.
        If omitted, a new numbering scheme will be started.
    level : int or None
        The level of the paragraph within the outline. If ``prev`` is
        set, defaults to the same level as in ``prev``. Otherwise,
        defaults to zero.
    num : bool
        If ``prev`` is :py:obj:`None` and the style of the paragraph
        does not correspond to an existing numbering style, this will
        determine wether or not the list will be numbered or bulleted.
        The result is not guaranteed, but is fairly safe for most Word
        templates.
    """
    xpath_options = {
        True: {'single': 'count(w:lvl)=1 and ', 'level': 0},
        False: {'single': '', 'level': level},
    }

    def style_xpath(prefer_single=True):
        """
        The style comes from the outer-scope variable ``par.style.name``.
        """
        style = par.style.style_id
        return (
            'w:abstractNum['
            '{single}w:lvl[@w:ilvl="{level}"]/w:pStyle[@w:val="{style}"]'
            ']/@w:abstractNumId'
        ).format(style=style, **xpath_options[prefer_single])

    def type_xpath(prefer_single=True):
        """
        The type is from the outer-scope variable ``num``.
        """
        type = 'decimal' if num else 'bullet'
        return (
            'w:abstractNum['
            '{single}w:lvl[@w:ilvl="{level}"]/w:numFmt[@w:val="{type}"]'
            ']/@w:abstractNumId'
        ).format(type=type, **xpath_options[prefer_single])

    def get_abstract_id():
        """
        Select as follows:

            1. Match single-level by style (get min ID)
            2. Match exact style and level (get min ID)
            3. Match single-level decimal/bullet types (get min ID)
            4. Match decimal/bullet in requested level (get min ID)
            3. 0
        """
        for fn in (style_xpath, type_xpath):
            for prefer_single in (True, False):
                xpath = fn(prefer_single)
                ids = numbering.xpath(xpath)
                if ids:
                    return min(int(x) for x in ids)
        return 0

    if (prev is None or
            prev._p.pPr is None or
            prev._p.pPr.numPr is None or
            prev._p.pPr.numPr.numId is None):
        if level is None:
            level = 0
        numbering = doc.part.numbering_part.numbering_definitions._numbering
        # Compute the abstract ID first by style, then by num
        anum = get_abstract_id()
        # Set the concrete numbering based on the abstract numbering ID
        num = numbering.add_num(anum)
        # Make sure to override the abstract continuation property
        num.add_lvlOverride(ilvl=level).add_startOverride(1)
        # Extract the newly-allocated concrete numbering ID
        num = num.numId
    else:
        if level is None:
            level = prev._p.pPr.numPr.ilvl.val
        # Get the previous concrete numbering ID
        num = prev._p.pPr.numPr.numId.val
    par._p.get_or_add_pPr().get_or_add_numPr().get_or_add_numId().val = num
    par._p.get_or_add_pPr().get_or_add_numPr().get_or_add_ilvl().val = level


from docx import Document

if __name__ == '__main__':
    # doc = Document('test2_1.docx')
    document = Document()
    para = document.add_paragraph()
    run = para.add_run(
        "Many font properties are tri-state, meaning they can take the values True, False, and None. True means "
        "theproperty is  False means it is . Conceptually, the None value means “inherit”. A run exists in the style "
        "hierarchy and by default inherits its character formatting from that hierarchy. Any character "
        "formattingdirectly applied using the Font object overrides the inherited values.")

    font = run.font
    font.name = 'Microsoft Yahei'
    font.size = Pt(10)

    format = para.paragraph_format
    # format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    format.first_line_indent = Pt(25)
    format.line_spacing = 2
    document.add_paragraph("item1", style="List 2")
    document.save("hch.docx")

# 页眉页脚
# len =len(doc.sections)
# print(len)
# content=doc.sections[0].footer.paragraphs
# for item in content:
#     print(item.text)


# 表格
# print(type(doc.tables))
# table=doc.tables[0]
# print(table.style,table.cell(0,0).text)
#
# print(doc.tables[0].rows[0].cells[0].text)


# print("共有" + str(len(doc.tables)) + "个表格")
# for i in range(len(doc.tables)):
#     table = doc.tables[i]
#     print("这是第" + str(i + 1) + "个表格：\n-----------------------------------")
#     for row in table.rows:
#         for cell in row.cells:
#             print("|" + cell.text, end="")
#         print("|\n")


# for style in styles:
#     print(style.priority,style.name)


# 每一段的内容

# for para in doc.paragraphs:
#     if len(para.runs) > 0:
#         size = para.runs[0].font.size
#         font_name = para.runs[0].font.name
#         if size != None:
#             print(para.text,"font_name：",font_name,"font_size：",size.pt)

# 每一段的编号、内容
# for i in range(len(doc.paragraphs)):
#     for j in range(len(doc.paragraphs[i].runs)):
#         print(str(i),'-',str(j), doc.paragraphs[i].runs[j].text)
